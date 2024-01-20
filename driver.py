import json
import docx
import random
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
Tense_Database = {
    "sai cấu trúc ngữ pháp": ["simple present", "present continuous", "present perfect", "present perfect continuous",
                          "simple past", "past continuous", "past perfect", "past perfect continuous",
                          "simple future", "future continuous", "future perfect", "future perfect continuous"],
}
with open('Database/database.json', 'r', encoding='utf-8') as file: Database = json.load(file)
with open('Database/wrong_grammatical_structure.json', 'r', encoding='utf-8') as file: Database_wrong_tense = json.load(file)

def Make_Docx(doc,Step:int,list_Data:list):
	for list_data in list_Data:
		Step += 1

		# Thêm câu hỏi
		doc.add_paragraph(f'Câu hỏi {Step}:').alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
		doc.add_paragraph(list_data['sentence']).alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

		# Thêm câu trả lời
		for option in list_data['options']:
		    doc.add_paragraph(f'- {option}').alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

		# Thêm đáp án
		doc.add_paragraph('Đáp án:').alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
		doc.add_paragraph(list_data['answer']).alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

def Number_of_Exercise(Data:list,Lose_Point:int,number_of_Exercise:float,location:str):
	number_of_exercise = []
	Percent_Point = Lose_Point / number_of_Exercise
	for data in Data:
		Exercise = data[2] // Percent_Point
		Exercise = int(Exercise)
		number_of_exercise.append(Exercise)

	doc = docx.Document()
	# Thêm tiêu đề
	doc.add_heading('Bài tập', level=1).alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
	step = 0
	Step = 0
	for data in  Data:
		last = ""
		num =  step +1
		num = str(num)
		number = number_of_exercise[step]
		step += 1

		text = data[0].upper()
		if data[1] !="*":
			last ="(" + str(data[1]).upper() +")"
		num += ") " + text + " "+last
		doc.add_paragraph(num).alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
		if data[1] == "*":	
			if number <  len(Database[data[0]]):
				list_Data = random.sample(Database[data[0]], k=number)
				Make_Docx(doc,Step,list_Data)		
			else:
				Make_Docx(Step,Database[data[0]])
		else:
			if number <  len(Database_wrong_tense[data[1]]):
				list_Data = random.sample(Database_wrong_tense[data[1]], k=number)
				Make_Docx(doc,Step,list_Data)		
			else:
				Make_Docx(Step,Database_wrong_tense[data[1]])
	# Lưu tài liệu Word xuống tệp
	name_save = location+ r"\cau_hoi.docx"
	doc.save(name_save)


def Analysis_word(Data:list,number_of_Exercise:float,location:str):
	Stored = ["sai mạo từ","sai trật tự từ","sai danh từ số ít-số nhiều","sai giới từ","sai các liên từ"]
	#Lose_Point tổng hợp số lượt sai của tất cả các lỗi
	Lose_Point = 0

	for data in Data:
		if data[0] in Stored:
			Lose_Point += data[2]
			data[1] = "*"
		else:
			if data[0] in Tense_Database:
				Lose_Point += data[2]
	Number_of_Exercise(Data,Lose_Point,number_of_Exercise,location)


def Handle_text(Link,number_of_Exercise:float,location:str):
	# Chia thành danh sách các dòng và loại bỏ dòng đầu
	with open(Link, "r", encoding="utf-8") as file:
	    raw_text = file.read()

	lines = raw_text.splitlines()
	lines = lines[1:]

	# Kết hợp danh sách các dòng thành văn bản mới
	new_text = '\n'.join(lines)

	Lines = new_text.split('\n')

	Handle_Data = []

	for line in Lines:
	    parts = line.split(',')
	    
	    if len(parts) == 3:
	        sub_list = [parts[0], parts[1], int(parts[2])]
	        Handle_Data.append(sub_list)
	
	Analysis_word(Handle_Data,number_of_Exercise,location)





